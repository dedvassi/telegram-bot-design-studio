import docx
import os

def extract_text_from_docx(file_path):
    """
    Extract text from a DOCX file.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text from the document
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text.append(paragraph.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error extracting text: {str(e)}"

if __name__ == "__main__":
    file_path = "/home/ubuntu/upload/ТЗ_Telegram_бот_дизайн_студия.docx"
    output_path = "/home/ubuntu/telegram_bot_project/requirements.txt"
    
    if os.path.exists(file_path):
        text = extract_text_from_docx(file_path)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
            
        print(f"Text extracted and saved to {output_path}")
    else:
        print(f"File not found: {file_path}")
